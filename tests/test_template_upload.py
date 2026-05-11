from __future__ import annotations

import asyncio
from datetime import datetime, timezone
from io import BytesIO
from uuid import uuid4

import pytest
from docx import Document
from fastapi import HTTPException, UploadFile

from app.enums import TemplateKind
from app.models.document_template import DocumentTemplate
from app.routers import templates as templates_router
from app.services.document_context import build_reference_render_context
from app.services.document_renderer import render_docx_bytes


def _docx_with_placeholder(text: str) -> bytes:
    doc = Document()
    doc.add_paragraph(text)
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


@pytest.mark.parametrize("kind", list(TemplateKind))
def test_reference_context_has_provider_and_address_keys(kind: TemplateKind) -> None:
    ctx = build_reference_render_context(kind)
    assert ctx["provider_full_name"]
    assert ctx["address_full"]
    assert ctx["provider_inn"]


def test_reference_context_for_initial_guarantee_uses_planned_name() -> None:
    ctx = build_reference_render_context(TemplateKind.GUARANTEE_INITIAL)
    assert ctx.get("client_planned_name")
    assert "client_full_name" not in ctx


def test_reference_context_for_contract_includes_pricing_and_term() -> None:
    ctx = build_reference_render_context(TemplateKind.CONTRACT)
    assert ctx["term_months"] in (6, 11)
    assert ctx["price_total_formatted"]
    assert ctx["price_total_in_words"]
    assert ctx["client_full_name"]


def test_reference_context_includes_correspondence_amount_when_paid() -> None:
    ctx = build_reference_render_context(TemplateKind.CONTRACT)
    assert ctx["has_correspondence_service"] is True
    assert ctx["correspondence_price_formatted"]
    assert ctx["correspondence_price_in_words"]


def test_render_docx_bytes_renders_simple_placeholder() -> None:
    template_bytes = _docx_with_placeholder("Привет, {{ provider_full_name }}!")
    result = render_docx_bytes(
        template_bytes=template_bytes,
        context={"provider_full_name": "ООО «Тест»"},
    )
    rendered_doc = Document(BytesIO(result))
    text = "\n".join(p.text for p in rendered_doc.paragraphs)
    assert "ООО «Тест»" in text
    assert "{{" not in text


def test_render_docx_bytes_raises_on_broken_jinja() -> None:
    template_bytes = _docx_with_placeholder("Сломано: {{ provider_full_name ")
    with pytest.raises(Exception):
        render_docx_bytes(template_bytes=template_bytes, context={"provider_full_name": "x"})


class _FakeUploadFile:
    """Минимальная замена fastapi.UploadFile — handler читает только filename и read()."""

    def __init__(self, *, filename: str, content: bytes):
        self.filename = filename
        self.content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        self._content = content

    async def read(self) -> bytes:
        return self._content


class _FakeDB:
    def __init__(self) -> None:
        self.added: list[object] = []
        self.committed = False

    async def execute(self, _stmt):
        class _R:
            def scalar_one(self_inner) -> int:
                return 0

        return _R()

    def add(self, item: object) -> None:
        self.added.append(item)

    async def commit(self) -> None:
        self.committed = True
        for item in self.added:
            if isinstance(item, DocumentTemplate) and item.id is None:
                item.id = uuid4()
                item.uploaded_at = datetime.now(timezone.utc)

    async def refresh(self, item: object) -> None:
        return None


def test_upload_handler_accepts_valid_template(tmp_path, monkeypatch) -> None:
    monkeypatch.setattr(
        "app.routers.templates.template_storage_dir",
        lambda kind: tmp_path / kind,
    )
    db = _FakeDB()
    template_bytes = _docx_with_placeholder("Договор для {{ provider_full_name }}")
    upload = _FakeUploadFile(filename="contract.docx", content=template_bytes)

    result = asyncio.get_event_loop().run_until_complete(
        templates_router.upload_template(
            kind=TemplateKind.CONTRACT, file=upload, comment=None, db=db
        )
    )

    assert result.test_render_succeeded is True
    saved = [item for item in db.added if isinstance(item, DocumentTemplate)]
    assert len(saved) == 1
    assert saved[0].kind == TemplateKind.CONTRACT.value
    assert (tmp_path / TemplateKind.CONTRACT.value).exists()
    files_on_disk = list((tmp_path / TemplateKind.CONTRACT.value).glob("*.docx"))
    assert len(files_on_disk) == 1


def test_upload_handler_rejects_broken_template(tmp_path, monkeypatch) -> None:
    monkeypatch.setattr(
        "app.routers.templates.template_storage_dir",
        lambda kind: tmp_path / kind,
    )
    db = _FakeDB()
    template_bytes = _docx_with_placeholder("Сломано {% if x")
    upload = _FakeUploadFile(filename="contract.docx", content=template_bytes)

    with pytest.raises(HTTPException) as exc_info:
        asyncio.get_event_loop().run_until_complete(
            templates_router.upload_template(
                kind=TemplateKind.CONTRACT, file=upload, comment=None, db=db
            )
        )

    assert exc_info.value.status_code == 422
    assert "Тестовый рендер" in exc_info.value.detail
    assert not [item for item in db.added if isinstance(item, DocumentTemplate)]
    target_dir = tmp_path / TemplateKind.CONTRACT.value
    if target_dir.exists():
        assert not list(target_dir.glob("*.docx"))
